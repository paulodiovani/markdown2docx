"""markdown2docx -- Convert GitHub Flavored Markdown files to DOCX format."""

from pathlib import Path

import click
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.image.image import Image as DocxImage
from docx.shared import Inches, Pt, RGBColor
from pygments import lex
from pygments.lexers import TextLexer, get_lexer_by_name
from pygments.token import (
    Comment,
    Error,
    Generic,
    Keyword,
    Literal,
    Name,
    Number,
    Operator,
    Punctuation,
    String,
    Token,
)

from lib.alerts import ALERT_STYLES, preprocess_alerts
from lib.mermaid import MERMAID_THEMES, preprocess_mermaid
from lib.parser import (
    create_parser,
    extract_text,
    heading_slug,
    preprocess_images,
    resolve_image_path,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

CODE_FONT = "Courier New"
CODE_FONT_SIZE = Pt(9)
CODE_BG_COLOR = "F2F2F2"
MAX_IMAGE_WIDTH = Inches(6)
MAX_IMAGE_HEIGHT = Inches(8)

_bookmark_id_counter = 0

# Pygments token type -> (hex_color, bold, italic)
TOKEN_STYLES = {
    Comment: ("008000", False, True),
    Comment.Preproc: ("AF00DB", True, False),
    Keyword: ("0000FF", True, False),
    Keyword.Constant: ("0000FF", True, False),
    Keyword.Type: ("267F99", False, False),
    Name.Builtin: ("795E26", False, False),
    Name.Function: ("795E26", False, False),
    Name.Class: ("267F99", True, False),
    Name.Decorator: ("795E26", False, False),
    Name.Exception: ("267F99", False, False),
    Name.Tag: ("800000", False, False),
    Name.Attribute: ("FF0000", False, False),
    String: ("A31515", False, False),
    Number: ("098658", False, False),
    Literal: ("A31515", False, False),
    Operator: ("000000", False, False),
    Operator.Word: ("0000FF", True, False),
    Punctuation: ("000000", False, False),
    Generic.Heading: ("0000FF", True, False),
    Generic.Subheading: ("0000FF", True, False),
    Generic.Emph: (None, False, True),
    Generic.Strong: (None, True, False),
    Generic.Error: ("FF0000", False, False),
    Error: ("FF0000", False, False),
}

# ---------------------------------------------------------------------------
# Pure helper functions
# ---------------------------------------------------------------------------


def get_token_style(token_type):
    """Walk Pygments token hierarchy to find matching style."""
    t = token_type
    while t is not Token:
        if t in TOKEN_STYLES:
            return TOKEN_STYLES[t]
        t = t.parent
    return None


def calculate_image_dimensions(img_path, max_width, max_height):
    """Calculate dimensions to fit within bounding box, preserving aspect ratio."""
    image = DocxImage.from_file(str(img_path))
    if image.px_width == 0 or image.px_height == 0:
        return max_width, max_height
    aspect_ratio = image.px_width / image.px_height
    width = max_width
    height = int(max_width / aspect_ratio)
    if height > max_height:
        height = max_height
        width = int(max_height * aspect_ratio)
    return width, height


# ---------------------------------------------------------------------------
# Side-effect functions (mutate doc or paragraph)
# ---------------------------------------------------------------------------


def set_paragraph_shading(paragraph, color_hex):
    """Set background shading on a paragraph via XML manipulation."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph via XML manipulation."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run_elem.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run_elem.append(text_elem)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph, anchor, text):
    """Add a hyperlink to a bookmark within the same document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run_elem.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run_elem.append(text_elem)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def add_image(doc, url, base_dir):
    """Resolve image path and add picture to document, scaled to fit page."""
    img_path = resolve_image_path(url, base_dir)
    if not img_path.exists():
        doc.add_paragraph(f"[Image not found: {url}]")
        return
    width, height = calculate_image_dimensions(
        img_path, MAX_IMAGE_WIDTH, MAX_IMAGE_HEIGHT
    )
    doc.add_picture(str(img_path), width=width, height=height)


# ---------------------------------------------------------------------------
# Rendering functions
# ---------------------------------------------------------------------------


def render_tokens(doc, tokens, base_dir):
    """Iterate top-level tokens and dispatch each."""
    for token in tokens:
        render_block(doc, token, base_dir)


def render_block(doc, token, base_dir):
    """Dispatch to render_{type} by token type."""
    t = token["type"]
    dispatch = {
        "heading": render_heading,
        "paragraph": render_paragraph,
        "block_code": render_block_code,
        "block_quote": render_block_quote,
        "alert": render_alert,
        "list": render_list,
        "table": render_table,
        "thematic_break": render_thematic_break,
    }
    handler = dispatch.get(t)
    if handler:
        if t in ("block_code", "thematic_break"):
            handler(doc, token)
        else:
            handler(doc, token, base_dir)


def render_heading(doc, token, base_dir):
    """Render a heading token."""
    global _bookmark_id_counter
    level = token.get("attrs", {}).get("level", 1) if token.get("attrs") else 1
    heading = doc.add_heading(level=level)
    render_inline(heading, token.get("children", []), base_dir)

    # Add a bookmark so internal links can target this heading
    text = extract_text(token.get("children", []))
    bookmark_name = heading_slug(text)
    _bookmark_id_counter += 1
    bid = str(_bookmark_id_counter)

    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), bid)
    bookmark_start.set(qn("w:name"), bookmark_name)
    heading._p.insert(0, bookmark_start)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), bid)
    heading._p.append(bookmark_end)


def render_paragraph(doc, token, base_dir):
    """Render a paragraph token with inline children."""
    children = token.get("children", [])

    # Check if paragraph contains only an image
    if len(children) == 1 and children[0]["type"] == "image":
        src = children[0].get("attrs", {}).get("src", "")
        if src:
            add_image(doc, src, base_dir)
        return

    para = doc.add_paragraph()
    render_inline(para, children, base_dir)


def render_block_code(doc, token):
    """Render a code block with Pygments syntax highlighting."""
    raw = token.get("raw", "") or token.get("text", "")
    info = token.get("attrs", {}).get("info", "") if token.get("attrs") else ""
    lang = info.split()[0] if info else ""

    # Strip trailing newline for cleaner rendering
    if raw.endswith("\n"):
        raw = raw[:-1]

    try:
        lexer = get_lexer_by_name(lang) if lang else TextLexer()
    except Exception:
        lexer = TextLexer()

    para = doc.add_paragraph()
    set_paragraph_shading(para, CODE_BG_COLOR)

    tokens_list = list(lex(raw, lexer))
    for token_type, value in tokens_list:
        if not value:
            continue
        run = para.add_run(value)
        run.font.name = CODE_FONT
        run.font.size = CODE_FONT_SIZE

        style = get_token_style(token_type)
        if style:
            color_hex, bold, italic = style
            if color_hex:
                run.font.color.rgb = RGBColor.from_string(color_hex)
            run.bold = bold
            run.italic = italic


def render_block_quote(doc, token, base_dir):
    """Render a blockquote with left indentation and gray left border."""
    children = token.get("children", [])
    for child in children:
        if child["type"] == "paragraph":
            para = doc.add_paragraph()

            # Add left indentation
            p_fmt = para.paragraph_format
            p_fmt.left_indent = Inches(0.5)

            # Add gray left border via XML
            p_pr = para._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            left_bdr = OxmlElement("w:left")
            left_bdr.set(qn("w:val"), "single")
            left_bdr.set(qn("w:sz"), "12")
            left_bdr.set(qn("w:space"), "4")
            left_bdr.set(qn("w:color"), "999999")
            p_bdr.append(left_bdr)
            p_pr.append(p_bdr)

            render_inline(para, child.get("children", []), base_dir)
        else:
            render_block(doc, child, base_dir)


def render_alert(doc, token, base_dir):
    """Render a GitHub-style alert as a colored box with label and body."""
    alert_type = token.get("attrs", {}).get("alert_type", "NOTE")
    border_color, bg_color, label_text, text_color = ALERT_STYLES[alert_type]
    children = token.get("children", [])

    def apply_alert_formatting(para):
        """Apply colored left border, background shading, and indent to a paragraph."""
        p_fmt = para.paragraph_format
        p_fmt.left_indent = Inches(0.5)

        set_paragraph_shading(para, bg_color)

        p_pr = para._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left_bdr = OxmlElement("w:left")
        left_bdr.set(qn("w:val"), "single")
        left_bdr.set(qn("w:sz"), "12")
        left_bdr.set(qn("w:space"), "4")
        left_bdr.set(qn("w:color"), border_color)
        p_bdr.append(left_bdr)
        p_pr.append(p_bdr)

    # Label paragraph
    label_para = doc.add_paragraph()
    apply_alert_formatting(label_para)
    run = label_para.add_run(label_text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(text_color)

    # Body paragraphs
    for child in children:
        if child["type"] == "paragraph":
            para = doc.add_paragraph()
            apply_alert_formatting(para)
            render_inline(para, child.get("children", []), base_dir)
        elif child["type"] == "blank_line":
            continue
        else:
            render_block(doc, child, base_dir)


def render_list(doc, token, base_dir):
    """Render ordered, unordered, and task lists."""
    attrs = token.get("attrs", {}) or {}
    ordered = attrs.get("ordered", False)
    children = token.get("children", [])

    for i, item in enumerate(children):
        item_children = item.get("children", [])

        # Determine style
        style = "List Number" if ordered else "List Bullet"

        # Check for task list item (type is task_list_item)
        is_task = item.get("type") == "task_list_item"
        checked = (
            item.get("attrs", {}).get("checked", False) if item.get("attrs") else False
        )

        for j, child in enumerate(item_children):
            if child["type"] in ("paragraph", "block_text"):
                para = doc.add_paragraph(style=style)
                if is_task and j == 0:
                    checkbox = "\u2611 " if checked else "\u2610 "
                    para.add_run(checkbox)
                render_inline(para, child.get("children", []), base_dir)
            elif child["type"] == "list":
                render_list(doc, child, base_dir)
            else:
                render_block(doc, child, base_dir)


def render_table(doc, token, base_dir):
    """Render a table with Table Grid style, bold headers, and alignment."""
    children = token.get("children", [])
    if not children:
        return

    # Find head and body -- table_head IS a row (contains cells directly),
    # table_body contains table_row children
    head = None
    body_rows = []
    for child in children:
        if child["type"] == "table_head":
            head = child
        elif child["type"] == "table_body":
            body_rows = child.get("children", [])

    # Build list of (cells_list, is_header) tuples
    all_rows = []
    if head:
        all_rows.append((head.get("children", []), True))
    for row in body_rows:
        all_rows.append((row.get("children", []), False))

    if not all_rows:
        return

    num_cols = len(all_rows[0][0])
    num_rows = len(all_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

    # Get column alignments from head cells
    aligns = []
    if head:
        for cell_token in head.get("children", []):
            cell_attrs = cell_token.get("attrs", {}) or {}
            aligns.append(cell_attrs.get("align"))

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    for row_idx, (cells_list, is_header) in enumerate(all_rows):
        for col_idx, cell_token in enumerate(cells_list):
            cell = table.rows[row_idx].cells[col_idx]
            para = cell.paragraphs[0]

            cell_children = cell_token.get("children", [])
            render_inline(para, cell_children, base_dir, bold=is_header)

            # Apply alignment
            if col_idx < len(aligns) and aligns[col_idx] in align_map:
                para.alignment = align_map[aligns[col_idx]]


def render_thematic_break(doc, token):
    """Render a horizontal rule as a paragraph with bottom border."""
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom_bdr = OxmlElement("w:bottom")
    bottom_bdr.set(qn("w:val"), "single")
    bottom_bdr.set(qn("w:sz"), "6")
    bottom_bdr.set(qn("w:space"), "1")
    bottom_bdr.set(qn("w:color"), "auto")
    p_bdr.append(bottom_bdr)
    p_pr.append(p_bdr)


def render_inline(
    paragraph, children, base_dir, bold=False, italic=False, strike=False
):
    """Recursive inline renderer for text, strong, emphasis, code, links, etc."""
    if children is None:
        return

    for child in children:
        t = child["type"]

        if t == "text":
            raw = (
                child.get("raw", "")
                or child.get("text", "")
                or child.get("children", "")
            )
            if isinstance(raw, list):
                raw = extract_text(raw)
            run = paragraph.add_run(raw)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if strike:
                run.font.strike = True

        elif t == "strong":
            render_inline(
                paragraph,
                child.get("children", []),
                base_dir,
                bold=True,
                italic=italic,
                strike=strike,
            )

        elif t == "emphasis":
            render_inline(
                paragraph,
                child.get("children", []),
                base_dir,
                bold=bold,
                italic=True,
                strike=strike,
            )

        elif t == "strikethrough":
            render_inline(
                paragraph,
                child.get("children", []),
                base_dir,
                bold=bold,
                italic=italic,
                strike=True,
            )

        elif t == "codespan":
            raw = (
                child.get("raw", "")
                or child.get("text", "")
                or child.get("children", "")
            )
            if isinstance(raw, list):
                raw = extract_text(raw)
            run = paragraph.add_run(raw)
            run.font.name = CODE_FONT
            run.font.size = CODE_FONT_SIZE
            if bold:
                run.bold = True
            if italic:
                run.italic = True

        elif t == "link":
            attrs = child.get("attrs", {}) or {}
            url = attrs.get("url", "") or attrs.get("href", "")
            link_text = extract_text(child.get("children", []))
            if url and url.startswith("#"):
                add_internal_hyperlink(paragraph, url[1:], link_text)
            elif url:
                add_hyperlink(paragraph, url, link_text)

        elif t == "image":
            src = child.get("attrs", {}).get("src", "")
            if src:
                add_image(paragraph._parent, src, base_dir)

        elif t == "softbreak":
            paragraph.add_run("\n")

        elif t == "linebreak":
            run = paragraph.add_run()
            run.add_break()


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------


def convert_file(input_path, output_dir, theme=None, transparent_bg=False):
    """Read MD, parse to AST, preprocess mermaid, render DOCX, save."""
    input_path = Path(input_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    md_text = input_path.read_text(encoding="utf-8")
    base_dir = str(input_path.parent)

    # Parse to AST with GFM plugins
    md = create_parser()
    tokens = md(md_text)

    # Preprocess mermaid diagrams
    tokens = preprocess_mermaid(
        tokens, base_dir, theme=theme, transparent_bg=transparent_bg
    )

    # Preprocess GitHub-style alerts
    tokens = preprocess_alerts(tokens)

    # Normalize image attrs (url -> src)
    tokens = preprocess_images(tokens)

    # Create document and render
    doc = Document()
    render_tokens(doc, tokens, base_dir)

    # Save
    output_path = output_dir / f"{input_path.name}.docx"
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


@click.command()
@click.argument("files", nargs=-1, required=True, type=click.Path(exists=True))
@click.option(
    "--output",
    "-o",
    default="./output",
    type=click.Path(),
    help="Output directory (default: ./output)",
)
@click.option(
    "--theme",
    type=click.Choice(MERMAID_THEMES),
    default=None,
    help="Mermaid diagram theme.",
)
@click.option(
    "--transparent",
    is_flag=True,
    default=False,
    help="Use transparent background for Mermaid diagrams.",
)
def main(files, output, theme, transparent):
    """Convert one or more Markdown files to DOCX format."""
    for f in files:
        out = convert_file(f, output, theme=theme, transparent_bg=transparent)
        click.echo(f"Converted: {f} -> {out}")


if __name__ == "__main__":
    main()
