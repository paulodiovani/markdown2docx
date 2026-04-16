"""Unit tests for markdown2docx.py (pure helpers + render_* functions)."""

from docx import Document
from docx.oxml.ns import qn
from pygments.token import Comment, Keyword, Name, Text

import markdown2docx as m2d
from markdown2docx import (
    ALERT_STYLES,
    CODE_BG_COLOR,
    CODE_FONT,
    add_hyperlink,
    add_internal_hyperlink,
    calculate_image_dimensions,
    get_token_style,
    render_alert,
    render_block_code,
    render_block_quote,
    render_heading,
    render_inline,
    render_list,
    render_paragraph,
    render_table,
    render_thematic_break,
    set_paragraph_shading,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _pPr(para):
    return para._p.find(qn("w:pPr"))


def _find_shd_fill(para):
    p_pr = _pPr(para)
    if p_pr is None:
        return None
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _find_left_border_color(para):
    p_pr = _pPr(para)
    if p_pr is None:
        return None
    bdr = p_pr.find(qn("w:pBdr"))
    if bdr is None:
        return None
    left = bdr.find(qn("w:left"))
    if left is None:
        return None
    return left.get(qn("w:color"))


def _find_bottom_border(para):
    p_pr = _pPr(para)
    if p_pr is None:
        return None
    bdr = p_pr.find(qn("w:pBdr"))
    if bdr is None:
        return None
    return bdr.find(qn("w:bottom"))


# ---------------------------------------------------------------------------
# get_token_style
# ---------------------------------------------------------------------------


def test_get_token_style_exact_match():
    style = get_token_style(Keyword)
    assert style == ("0000FF", True, False)


def test_get_token_style_walks_to_parent():
    # Name.Variable has no entry; should walk up but eventually return None
    # Name.Function has an entry; confirm we find it via Name.Function itself
    style = get_token_style(Name.Function)
    assert style == ("795E26", False, False)


def test_get_token_style_returns_none_for_unmapped_root():
    assert get_token_style(Text) is None


def test_get_token_style_inherits_from_parent_token():
    # Comment.Multiline has no direct entry; should inherit from Comment
    style = get_token_style(Comment.Multiline)
    assert style == ("008000", False, True)


# ---------------------------------------------------------------------------
# calculate_image_dimensions
# ---------------------------------------------------------------------------


class _FakeImage:
    """Stand-in for ``docx.image.image.Image`` with fixed px dimensions."""

    def __init__(self, px_width, px_height):
        self.px_width = px_width
        self.px_height = px_height


def _stub_docx_image(monkeypatch, px_width, px_height):
    fake = _FakeImage(px_width, px_height)
    monkeypatch.setattr("markdown2docx.DocxImage.from_file", lambda _path: fake)


def test_calculate_image_dimensions_square_fits_width(monkeypatch):
    from docx.shared import Inches

    _stub_docx_image(monkeypatch, 100, 100)
    max_w = Inches(6)
    max_h = Inches(8)
    w, h = calculate_image_dimensions("ignored", max_w, max_h)
    assert w == max_w
    assert h == max_w  # square: height also = max_w


def test_calculate_image_dimensions_landscape_image(monkeypatch):
    from docx.shared import Inches

    _stub_docx_image(monkeypatch, 200, 50)  # aspect = 4
    max_w = Inches(6)
    max_h = Inches(8)
    w, h = calculate_image_dimensions("ignored", max_w, max_h)
    assert w == max_w
    assert h == int(max_w / 4)


def test_calculate_image_dimensions_tall_image_clamps_to_max_height(monkeypatch):
    from docx.shared import Inches

    _stub_docx_image(monkeypatch, 10, 200)  # aspect = 0.05
    max_w = Inches(6)
    max_h = Inches(4)
    w, h = calculate_image_dimensions("ignored", max_w, max_h)
    assert h == max_h
    assert w == int(max_h * (10 / 200))


def test_calculate_image_dimensions_zero_dim_returns_maxes(monkeypatch):
    from docx.shared import Inches

    _stub_docx_image(monkeypatch, 0, 0)
    max_w = Inches(6)
    max_h = Inches(8)
    w, h = calculate_image_dimensions("ignored", max_w, max_h)
    assert w == max_w
    assert h == max_h


# ---------------------------------------------------------------------------
# set_paragraph_shading
# ---------------------------------------------------------------------------


def test_set_paragraph_shading_adds_shd_element(make_doc):
    doc = make_doc()
    para = doc.add_paragraph("x")
    set_paragraph_shading(para, "ABCDEF")
    assert _find_shd_fill(para) == "ABCDEF"


# ---------------------------------------------------------------------------
# add_hyperlink / add_internal_hyperlink
# ---------------------------------------------------------------------------


def test_add_hyperlink_creates_hyperlink_element(make_doc):
    doc = make_doc()
    para = doc.add_paragraph()
    add_hyperlink(para, "https://example.com", "click me")
    hyperlinks = para._p.findall(qn("w:hyperlink"))
    assert len(hyperlinks) == 1
    assert hyperlinks[0].get(qn("r:id"))


def test_add_hyperlink_text_is_rendered(make_doc):
    doc = make_doc()
    para = doc.add_paragraph()
    add_hyperlink(para, "https://example.com", "click me")
    # Text appears inside the hyperlink element's run
    hl = para._p.find(qn("w:hyperlink"))
    t_elem = hl.find(".//" + qn("w:t"))
    assert t_elem.text == "click me"


def test_add_internal_hyperlink_uses_anchor_attr(make_doc):
    doc = make_doc()
    para = doc.add_paragraph()
    add_internal_hyperlink(para, "my-section", "see section")
    hl = para._p.find(qn("w:hyperlink"))
    assert hl.get(qn("w:anchor")) == "my-section"
    assert hl.get(qn("r:id")) is None


# ---------------------------------------------------------------------------
# render_heading
# ---------------------------------------------------------------------------


def test_render_heading_adds_heading_with_level(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "heading",
        "attrs": {"level": 2},
        "children": [{"type": "text", "raw": "Section Title"}],
    }
    render_heading(doc, token, str(tmp_path))
    last = doc.paragraphs[-1]
    assert "Heading 2" in last.style.name
    assert "Section Title" in last.text


def test_render_heading_inserts_bookmark(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "heading",
        "attrs": {"level": 1},
        "children": [{"type": "text", "raw": "Hello World"}],
    }
    render_heading(doc, token, str(tmp_path))
    last_p = doc.paragraphs[-1]._p
    bms = last_p.find(qn("w:bookmarkStart"))
    bme = last_p.find(qn("w:bookmarkEnd"))
    assert bms is not None and bme is not None
    assert bms.get(qn("w:name")) == "hello-world"


def test_render_heading_defaults_to_level_1_when_no_attrs(make_doc, tmp_path):
    doc = make_doc()
    token = {"type": "heading", "children": [{"type": "text", "raw": "H"}]}
    render_heading(doc, token, str(tmp_path))
    assert "Heading 1" in doc.paragraphs[-1].style.name


# ---------------------------------------------------------------------------
# render_paragraph
# ---------------------------------------------------------------------------


def test_render_paragraph_adds_text(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "paragraph",
        "children": [{"type": "text", "raw": "Hello"}],
    }
    render_paragraph(doc, token, str(tmp_path))
    assert doc.paragraphs[-1].text == "Hello"


def test_render_paragraph_image_only_emits_picture(make_doc, tmp_path, small_jpeg):
    doc = make_doc()
    token = {
        "type": "paragraph",
        "children": [
            {"type": "image", "attrs": {"src": small_jpeg.name}},
        ],
    }
    render_paragraph(doc, token, str(small_jpeg.parent))
    # An image added via add_picture appears as a paragraph with an inline drawing
    drawings = doc.element.body.findall(".//" + qn("w:drawing"))
    assert len(drawings) == 1


def test_render_paragraph_missing_image_adds_placeholder(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "paragraph",
        "children": [{"type": "image", "attrs": {"src": "nope.png"}}],
    }
    render_paragraph(doc, token, str(tmp_path))
    assert "[Image not found" in doc.paragraphs[-1].text


# ---------------------------------------------------------------------------
# render_block_code
# ---------------------------------------------------------------------------


def test_render_block_code_uses_code_font(make_doc):
    doc = make_doc()
    token = {"type": "block_code", "raw": "x = 1\n", "attrs": {"info": "python"}}
    render_block_code(doc, token)
    para = doc.paragraphs[-1]
    runs = para.runs
    assert runs
    assert all(r.font.name == CODE_FONT for r in runs if r.text)


def test_render_block_code_applies_shading(make_doc):
    doc = make_doc()
    token = {"type": "block_code", "raw": "x = 1", "attrs": {"info": "python"}}
    render_block_code(doc, token)
    assert _find_shd_fill(doc.paragraphs[-1]) == CODE_BG_COLOR


def test_render_block_code_no_lang_falls_back_to_text_lexer(make_doc):
    doc = make_doc()
    token = {"type": "block_code", "raw": "plain text", "attrs": {}}
    render_block_code(doc, token)
    # no exception, renders as one run
    runs = doc.paragraphs[-1].runs
    assert any(r.text for r in runs)


def test_render_block_code_unknown_lang_falls_back(make_doc):
    doc = make_doc()
    token = {
        "type": "block_code",
        "raw": "whatever",
        "attrs": {"info": "not-a-real-lang"},
    }
    render_block_code(doc, token)
    assert doc.paragraphs[-1].runs


def test_render_block_code_strips_trailing_newline(make_doc):
    """The renderer strips one trailing newline before lexing.

    Pygments' Python lexer adds a final ``\n`` to its token stream, so the
    rendered paragraph text will contain exactly one trailing newline — not
    two, which is what we'd see without the strip.
    """
    doc = make_doc()
    token = {"type": "block_code", "raw": "print('hi')\n", "attrs": {"info": "python"}}
    render_block_code(doc, token)
    text = doc.paragraphs[-1].text
    assert not text.endswith("\n\n")


# ---------------------------------------------------------------------------
# render_block_quote
# ---------------------------------------------------------------------------


def test_render_block_quote_adds_left_border(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "block_quote",
        "children": [
            {"type": "paragraph", "children": [{"type": "text", "raw": "quoted"}]}
        ],
    }
    render_block_quote(doc, token, str(tmp_path))
    assert _find_left_border_color(doc.paragraphs[-1]) == "999999"


def test_render_block_quote_indents_paragraph(make_doc, tmp_path):
    from docx.shared import Inches

    doc = make_doc()
    token = {
        "type": "block_quote",
        "children": [
            {"type": "paragraph", "children": [{"type": "text", "raw": "quoted"}]}
        ],
    }
    render_block_quote(doc, token, str(tmp_path))
    assert doc.paragraphs[-1].paragraph_format.left_indent == Inches(0.5)


# ---------------------------------------------------------------------------
# render_alert
# ---------------------------------------------------------------------------


def test_render_alert_emits_label_and_body(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "alert",
        "attrs": {"alert_type": "NOTE"},
        "children": [
            {"type": "paragraph", "children": [{"type": "text", "raw": "body"}]}
        ],
    }
    before = len(doc.paragraphs)
    render_alert(doc, token, str(tmp_path))
    new_paras = doc.paragraphs[before:]
    # one label + one body
    assert len(new_paras) >= 2
    assert new_paras[0].runs[0].bold is True
    _, _, expected_label, _ = ALERT_STYLES["NOTE"]
    assert new_paras[0].runs[0].text == expected_label


def test_render_alert_applies_correct_shading_per_type(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "alert",
        "attrs": {"alert_type": "WARNING"},
        "children": [{"type": "paragraph", "children": [{"type": "text", "raw": "x"}]}],
    }
    render_alert(doc, token, str(tmp_path))
    _, bg_color, _, _ = ALERT_STYLES["WARNING"]
    assert _find_shd_fill(doc.paragraphs[-1]) == bg_color


def test_render_alert_skips_blank_lines_in_body(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "alert",
        "attrs": {"alert_type": "TIP"},
        "children": [
            {"type": "blank_line"},
            {"type": "paragraph", "children": [{"type": "text", "raw": "body"}]},
        ],
    }
    before = len(doc.paragraphs)
    render_alert(doc, token, str(tmp_path))
    new_paras = doc.paragraphs[before:]
    # label + body (blank_line skipped)
    assert len(new_paras) == 2


# ---------------------------------------------------------------------------
# render_list
# ---------------------------------------------------------------------------


def test_render_list_unordered_applies_bullet_numbering(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "list",
        "attrs": {"ordered": False},
        "children": [
            {
                "type": "list_item",
                "children": [
                    {"type": "block_text", "children": [{"type": "text", "raw": "a"}]}
                ],
            }
        ],
    }
    before = len(doc.paragraphs)
    render_list(doc, token, str(tmp_path))
    new_para = doc.paragraphs[before]
    p_pr = _pPr(new_para)
    num_pr = p_pr.find(qn("w:numPr"))
    assert num_pr is not None


def test_render_list_ordered_uses_fresh_numid(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "list",
        "attrs": {"ordered": True},
        "children": [
            {
                "type": "list_item",
                "children": [
                    {"type": "block_text", "children": [{"type": "text", "raw": "1"}]}
                ],
            }
        ],
    }
    render_list(doc, token, str(tmp_path))
    p_pr = _pPr(doc.paragraphs[-1])
    num_pr = p_pr.find(qn("w:numPr"))
    num = num_pr.find(qn("w:numId"))
    assert num is not None


def test_render_list_task_list_item_adds_checkbox(make_doc, tmp_path):
    doc = make_doc()
    token = {
        "type": "list",
        "attrs": {"ordered": False},
        "children": [
            {
                "type": "task_list_item",
                "attrs": {"checked": True},
                "children": [
                    {
                        "type": "block_text",
                        "children": [{"type": "text", "raw": "done"}],
                    }
                ],
            },
            {
                "type": "task_list_item",
                "attrs": {"checked": False},
                "children": [
                    {
                        "type": "block_text",
                        "children": [{"type": "text", "raw": "todo"}],
                    }
                ],
            },
        ],
    }
    before = len(doc.paragraphs)
    render_list(doc, token, str(tmp_path))
    new_paras = doc.paragraphs[before:]
    assert new_paras[0].text.startswith("\u2611 ")  # checked
    assert new_paras[1].text.startswith("\u2610 ")  # unchecked


def test_render_list_separate_lists_get_distinct_numids(make_doc, tmp_path):
    doc = make_doc()
    token1 = {
        "type": "list",
        "attrs": {"ordered": True},
        "children": [
            {
                "type": "list_item",
                "children": [
                    {"type": "block_text", "children": [{"type": "text", "raw": "a"}]}
                ],
            }
        ],
    }
    token2 = {
        "type": "list",
        "attrs": {"ordered": True},
        "children": [
            {
                "type": "list_item",
                "children": [
                    {"type": "block_text", "children": [{"type": "text", "raw": "x"}]}
                ],
            }
        ],
    }
    render_list(doc, token1, str(tmp_path))
    p1 = doc.paragraphs[-1]
    render_list(doc, token2, str(tmp_path))
    p2 = doc.paragraphs[-1]

    def num_id(para):
        num_pr = _pPr(para).find(qn("w:numPr"))
        return num_pr.find(qn("w:numId")).get(qn("w:val"))

    assert num_id(p1) != num_id(p2)


# ---------------------------------------------------------------------------
# render_table
# ---------------------------------------------------------------------------


def _table_token():
    return {
        "type": "table",
        "children": [
            {
                "type": "table_head",
                "children": [
                    {
                        "type": "table_cell",
                        "attrs": {"align": "left"},
                        "children": [{"type": "text", "raw": "A"}],
                    },
                    {
                        "type": "table_cell",
                        "attrs": {"align": "right"},
                        "children": [{"type": "text", "raw": "B"}],
                    },
                ],
            },
            {
                "type": "table_body",
                "children": [
                    {
                        "type": "table_row",
                        "children": [
                            {
                                "type": "table_cell",
                                "children": [{"type": "text", "raw": "1"}],
                            },
                            {
                                "type": "table_cell",
                                "children": [{"type": "text", "raw": "2"}],
                            },
                        ],
                    }
                ],
            },
        ],
    }


def test_render_table_row_and_col_count(make_doc, tmp_path):
    doc = make_doc()
    render_table(doc, _table_token(), str(tmp_path))
    table = doc.tables[-1]
    assert len(table.rows) == 2
    assert len(table.columns) == 2


def test_render_table_header_cells_are_bold(make_doc, tmp_path):
    doc = make_doc()
    render_table(doc, _table_token(), str(tmp_path))
    header_para = doc.tables[-1].rows[0].cells[0].paragraphs[0]
    assert any(r.bold for r in header_para.runs)


def test_render_table_cell_alignment_applied(make_doc, tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = make_doc()
    render_table(doc, _table_token(), str(tmp_path))
    table = doc.tables[-1]
    assert table.rows[0].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_render_table_empty_children_returns_early(make_doc, tmp_path):
    doc = make_doc()
    render_table(doc, {"type": "table", "children": []}, str(tmp_path))
    assert doc.tables == []


# ---------------------------------------------------------------------------
# render_thematic_break
# ---------------------------------------------------------------------------


def test_render_thematic_break_adds_bottom_border(make_doc):
    doc = make_doc()
    render_thematic_break(doc, {"type": "thematic_break"})
    bottom = _find_bottom_border(doc.paragraphs[-1])
    assert bottom is not None
    assert bottom.get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# render_inline
# ---------------------------------------------------------------------------


def test_render_inline_strong_sets_bold(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [{"type": "strong", "children": [{"type": "text", "raw": "bold text"}]}]
    render_inline(para, children, str(tmp_path))
    assert para.runs[0].text == "bold text"
    assert para.runs[0].bold is True


def test_render_inline_emphasis_sets_italic(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [{"type": "emphasis", "children": [{"type": "text", "raw": "italic"}]}]
    render_inline(para, children, str(tmp_path))
    assert para.runs[0].italic is True


def test_render_inline_strikethrough_sets_strike(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [
        {
            "type": "strikethrough",
            "children": [{"type": "text", "raw": "struck"}],
        }
    ]
    render_inline(para, children, str(tmp_path))
    assert para.runs[0].font.strike is True


def test_render_inline_codespan_uses_code_font(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [{"type": "codespan", "raw": "x"}]
    render_inline(para, children, str(tmp_path))
    assert para.runs[0].font.name == CODE_FONT


def test_render_inline_external_link_creates_hyperlink(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [
        {
            "type": "link",
            "attrs": {"url": "https://example.com"},
            "children": [{"type": "text", "raw": "click"}],
        }
    ]
    render_inline(para, children, str(tmp_path))
    hl = para._p.find(qn("w:hyperlink"))
    assert hl is not None
    assert hl.get(qn("r:id")) is not None


def test_render_inline_internal_link_uses_anchor(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [
        {
            "type": "link",
            "attrs": {"url": "#target"},
            "children": [{"type": "text", "raw": "go"}],
        }
    ]
    render_inline(para, children, str(tmp_path))
    hl = para._p.find(qn("w:hyperlink"))
    assert hl is not None
    assert hl.get(qn("w:anchor")) == "target"


def test_render_inline_softbreak_adds_newline(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [
        {"type": "text", "raw": "a"},
        {"type": "softbreak"},
        {"type": "text", "raw": "b"},
    ]
    render_inline(para, children, str(tmp_path))
    assert "\n" in "".join(r.text for r in para.runs)


def test_render_inline_linebreak_adds_break_element(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [{"type": "linebreak"}]
    render_inline(para, children, str(tmp_path))
    assert para._p.find(".//" + qn("w:br")) is not None


def test_render_inline_none_children_returns_quietly(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    # Must not raise
    render_inline(para, None, str(tmp_path))
    assert para.runs == []


def test_render_inline_nested_strong_in_emphasis(make_doc, tmp_path):
    doc = make_doc()
    para = doc.add_paragraph()
    children = [
        {
            "type": "emphasis",
            "children": [
                {"type": "strong", "children": [{"type": "text", "raw": "both"}]}
            ],
        }
    ]
    render_inline(para, children, str(tmp_path))
    assert para.runs[0].bold is True
    assert para.runs[0].italic is True


# ---------------------------------------------------------------------------
# convert_file (thin integration — verifies orchestrator wiring)
# ---------------------------------------------------------------------------


def test_convert_file_creates_output_file(tmp_path):
    md = tmp_path / "sample.md"
    md.write_text("# Title\n\nHello, world.\n")
    out_dir = tmp_path / "out"
    result = m2d.convert_file(str(md), str(out_dir))
    assert result.exists()
    assert result.name == "sample.md.docx"


def test_convert_file_creates_missing_output_dir(tmp_path):
    md = tmp_path / "x.md"
    md.write_text("# H\n")
    out_dir = tmp_path / "nested" / "output"
    m2d.convert_file(str(md), str(out_dir))
    assert out_dir.is_dir()


def test_convert_file_result_is_valid_docx(tmp_path):
    md = tmp_path / "x.md"
    md.write_text("# Heading\n\nPara.\n")
    out_dir = tmp_path / "out"
    result = m2d.convert_file(str(md), str(out_dir))
    # Reopen the file to prove it's valid
    doc = Document(str(result))
    texts = [p.text for p in doc.paragraphs]
    assert any("Heading" in t for t in texts)
    assert any("Para" in t for t in texts)
