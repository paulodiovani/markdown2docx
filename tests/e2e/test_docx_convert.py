"""End-to-end tests for the DOCX conversion pipeline.

These tests run ``markdown2docx.convert_file()`` against real files in
``examples/`` and reopen the resulting ``.docx`` to assert its structure.
"""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

import markdown2docx as m2d


EXAMPLES = Path(__file__).resolve().parent.parent.parent / "examples"


@pytest.fixture
def no_mermaid(monkeypatch):
    """Stub ``subprocess.run`` so ``mmdc`` is never invoked.

    Touches the expected PNG output path so ``render_paragraph`` sees the
    file as existing (Mermaid blocks without a rendered PNG still reach the
    renderer, but the resulting image won't be valid — that's fine since
    python-docx's ``add_picture`` is what would fail, not our code. To keep
    this simple we copy ``examples/cat.jpg`` to the expected output path.
    """
    import shutil
    import subprocess

    cat = EXAMPLES / "cat.jpg"

    def fake_run(cmd, **kwargs):
        # The Mermaid preprocessor builds commands like:
        # ["mmdc", "-i", "<in>.mmd", "-o", "<out>.png", ...]
        if "-o" in cmd:
            out_path = cmd[cmd.index("-o") + 1]
            shutil.copy(cat, out_path)

    monkeypatch.setattr(subprocess, "run", fake_run)


def _get_texts(doc):
    return [p.text for p in doc.paragraphs]


# ---------------------------------------------------------------------------
# basic.md — all the bread-and-butter block types
# ---------------------------------------------------------------------------


def test_convert_basic_md_produces_valid_docx(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    assert result.exists()
    assert result.name == "basic.md.docx"


def test_convert_basic_md_has_all_heading_levels(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    doc = Document(str(result))
    levels = {
        p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")
    }
    assert "Heading 1" in levels
    assert "Heading 2" in levels
    assert "Heading 3" in levels


def test_convert_basic_md_emits_table(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    doc = Document(str(result))
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert len(tbl.rows) == 4  # 1 header + 3 body
    assert len(tbl.columns) == 3


def test_convert_basic_md_emits_code_blocks_with_shading(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    doc = Document(str(result))
    # at least one paragraph has a shd fill = code bg color
    fills = []
    for p in doc.paragraphs:
        p_pr = p._p.find(qn("w:pPr"))
        if p_pr is None:
            continue
        shd = p_pr.find(qn("w:shd"))
        if shd is not None:
            fills.append(shd.get(qn("w:fill")))
    assert "F2F2F2" in fills


def test_convert_basic_md_emits_hyperlink(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    doc = Document(str(result))
    hyperlinks = doc.element.body.findall(".//" + qn("w:hyperlink"))
    assert hyperlinks


def test_convert_basic_md_emits_thematic_break(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    doc = Document(str(result))
    # Find a paragraph with a bottom border
    found = False
    for p in doc.paragraphs:
        p_pr = p._p.find(qn("w:pPr"))
        if p_pr is None:
            continue
        bdr = p_pr.find(qn("w:pBdr"))
        if bdr is not None and bdr.find(qn("w:bottom")) is not None:
            found = True
            break
    assert found


def test_convert_basic_md_emits_task_list_with_checkboxes(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    doc = Document(str(result))
    all_text = "\n".join(_get_texts(doc))
    assert "\u2611" in all_text  # checked
    assert "\u2610" in all_text  # unchecked


# ---------------------------------------------------------------------------
# alerts.md — all five GFM alert types
# ---------------------------------------------------------------------------


def test_convert_alerts_md_renders_all_five_types(tmp_path):
    from markdown2docx import ALERT_STYLES

    result = m2d.convert_file(str(EXAMPLES / "alerts.md"), tmp_path)
    doc = Document(str(result))
    all_text = "\n".join(_get_texts(doc))

    for alert_type in ALERT_STYLES:
        _, _, label, _ = ALERT_STYLES[alert_type]
        assert label in all_text, f"missing alert label: {label}"


def test_convert_alerts_md_applies_shading_per_type(tmp_path):
    from markdown2docx import ALERT_STYLES

    result = m2d.convert_file(str(EXAMPLES / "alerts.md"), tmp_path)
    doc = Document(str(result))

    fills = set()
    for p in doc.paragraphs:
        p_pr = p._p.find(qn("w:pPr"))
        if p_pr is None:
            continue
        shd = p_pr.find(qn("w:shd"))
        if shd is not None:
            fills.add(shd.get(qn("w:fill")))

    expected_bgs = {bg for (_, bg, _, _) in ALERT_STYLES.values()}
    # every alert-type background should appear in the document
    assert expected_bgs.issubset(fills)


# ---------------------------------------------------------------------------
# full.md — mixed headings 1–6 + bookmarks + image
# ---------------------------------------------------------------------------


def test_convert_full_md_creates_bookmarks(tmp_path, no_mermaid):
    result = m2d.convert_file(str(EXAMPLES / "full.md"), tmp_path)
    doc = Document(str(result))
    bms = doc.element.body.findall(".//" + qn("w:bookmarkStart"))
    assert len(bms) >= 6  # at least one per heading in full.md


def test_convert_full_md_embeds_image(tmp_path, no_mermaid):
    result = m2d.convert_file(str(EXAMPLES / "full.md"), tmp_path)
    doc = Document(str(result))
    drawings = doc.element.body.findall(".//" + qn("w:drawing"))
    assert drawings, "expected at least one embedded image"


# ---------------------------------------------------------------------------
# internal-links.md — anchor-based hyperlinks
# ---------------------------------------------------------------------------


def test_convert_internal_links_md_uses_anchor_hyperlinks(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "internal-links.md"), tmp_path)
    doc = Document(str(result))
    hyperlinks = doc.element.body.findall(".//" + qn("w:hyperlink"))
    # At least some hyperlinks in this doc target internal anchors (w:anchor),
    # not external URLs (r:id).
    anchor_links = [hl for hl in hyperlinks if hl.get(qn("w:anchor"))]
    assert anchor_links


# ---------------------------------------------------------------------------
# mermaid.md — mmdc is mocked, PNG path stubbed
# ---------------------------------------------------------------------------


def test_convert_mermaid_md_invokes_mmdc_and_embeds_image(tmp_path, no_mermaid):
    result = m2d.convert_file(str(EXAMPLES / "mermaid.md"), tmp_path)
    doc = Document(str(result))
    drawings = doc.element.body.findall(".//" + qn("w:drawing"))
    # mermaid.md has 2 diagrams -> 2 drawings expected
    assert len(drawings) >= 2


# ---------------------------------------------------------------------------
# Output directory creation
# ---------------------------------------------------------------------------


def test_convert_creates_missing_output_directory(tmp_path):
    out = tmp_path / "a" / "b" / "c"
    m2d.convert_file(str(EXAMPLES / "basic.md"), out)
    assert out.is_dir()


def test_convert_writes_docx_named_after_source(tmp_path):
    result = m2d.convert_file(str(EXAMPLES / "basic.md"), tmp_path)
    assert result == tmp_path / "basic.md.docx"
